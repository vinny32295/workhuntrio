#!/usr/bin/env python3
"""
Job Hunter Runner - Combines discovery with application processing

This script:
1. Discovers new jobs via Google Search
2. Scrapes job details from discovered URLs
3. Analyzes fit with your profile
4. Generates tailored resumes and cover letters
5. Tracks everything in SQLite

Usage:
    # Full discovery + processing
    python run_hunter.py
    
    # Discovery only (find URLs, don't process)
    python run_hunter.py --discover-only
    
    # Process a single URL
    python run_hunter.py --url "https://boards.greenhouse.io/company/jobs/123"
    
    # Daemon mode (runs every hour)
    python run_hunter.py --daemon

Environment Variables:
    ANTHROPIC_API_KEY  - Required for AI analysis/tailoring
    GOOGLE_API_KEY     - For Google Custom Search (recommended)
    GOOGLE_CSE_ID      - Custom Search Engine ID
    SERPAPI_KEY        - Alternative to Google CSE
"""

import os
import sys
import json
import time
import hashlib
import argparse
import logging
import sqlite3
from datetime import datetime
from pathlib import Path
from dataclasses import dataclass, asdict
from typing import Optional

# Local imports
from google_discovery import JobDiscovery, URLClassifier, DEFAULT_SEARCH_QUERIES

# Third-party
import requests
from bs4 import BeautifulSoup
import anthropic
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import schedule

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('job_hunter.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Configuration
CONFIG = {
    "anthropic_api_key": os.getenv("ANTHROPIC_API_KEY", ""),
    "output_dir": Path("./output"),
    "db_path": Path("./jobs.db"),
    "master_resume_path": Path("./master_resume.json"),
    "max_jobs_per_run": 50,
    "min_match_score": 0.5,
    "target_roles": [
        "product manager", "operations manager", "project manager",
        "revenue operations", "sales operations", "business operations",
        "customer success", "program manager",
    ],
    "exclude_keywords": ["senior director", "vp", "vice president", "chief", "intern"],
}


@dataclass
class Job:
    id: str
    title: str
    company: str
    url: str
    location: str
    description: str
    requirements: list
    keywords: list
    discovered_at: str
    match_score: float
    source: str = "unknown"
    status: str = "new"


@dataclass
class MasterResume:
    name: str
    email: str
    phone: str
    location: str
    summary: str
    experience: list
    education: list
    skills: list
    certifications: list


# =============================================================================
# DATABASE
# =============================================================================

def init_db() -> sqlite3.Connection:
    conn = sqlite3.connect(CONFIG["db_path"])
    cursor = conn.cursor()
    
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS jobs (
            id TEXT PRIMARY KEY,
            title TEXT,
            company TEXT,
            url TEXT UNIQUE,
            location TEXT,
            description TEXT,
            requirements TEXT,
            keywords TEXT,
            discovered_at TEXT,
            match_score REAL,
            source TEXT,
            status TEXT DEFAULT 'new',
            resume_path TEXT,
            cover_letter_path TEXT,
            applied_at TEXT
        )
    """)
    
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS discovered_urls (
            url TEXT PRIMARY KEY,
            discovered_at TEXT,
            classification TEXT,
            processed INTEGER DEFAULT 0
        )
    """)
    
    conn.commit()
    return conn


def url_seen(conn: sqlite3.Connection, url: str) -> bool:
    cursor = conn.cursor()
    cursor.execute("SELECT 1 FROM discovered_urls WHERE url = ?", (url,))
    return cursor.fetchone() is not None


def save_discovered(conn: sqlite3.Connection, url: str, classification: dict):
    cursor = conn.cursor()
    cursor.execute("""
        INSERT OR IGNORE INTO discovered_urls (url, discovered_at, classification)
        VALUES (?, ?, ?)
    """, (url, datetime.now().isoformat(), json.dumps(classification)))
    conn.commit()


def save_job(conn: sqlite3.Connection, job: Job, resume_path: str = None, cl_path: str = None):
    cursor = conn.cursor()
    cursor.execute("""
        INSERT OR REPLACE INTO jobs 
        (id, title, company, url, location, description, requirements, keywords,
         discovered_at, match_score, source, status, resume_path, cover_letter_path)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        job.id, job.title, job.company, job.url, job.location, job.description,
        json.dumps(job.requirements), json.dumps(job.keywords),
        job.discovered_at, job.match_score, job.source, job.status,
        resume_path, cl_path
    ))
    conn.commit()


def get_pending_urls(conn: sqlite3.Connection, limit: int = 50) -> list[dict]:
    cursor = conn.cursor()
    cursor.execute("""
        SELECT url, classification FROM discovered_urls 
        WHERE processed = 0 
        ORDER BY discovered_at DESC 
        LIMIT ?
    """, (limit,))
    return [{'url': row[0], 'classification': json.loads(row[1])} for row in cursor.fetchall()]


def mark_processed(conn: sqlite3.Connection, url: str):
    cursor = conn.cursor()
    cursor.execute("UPDATE discovered_urls SET processed = 1 WHERE url = ?", (url,))
    conn.commit()


# =============================================================================
# SCRAPERS
# =============================================================================

class Scraper:
    def __init__(self):
        self.session = requests.Session()
        self.session.headers.update({
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
        })
    
    def get_page(self, url: str) -> Optional[BeautifulSoup]:
        try:
            response = self.session.get(url, timeout=30)
            response.raise_for_status()
            return BeautifulSoup(response.text, 'lxml')
        except Exception as e:
            logger.error(f"Failed to fetch {url}: {e}")
            return None
    
    def get_greenhouse_job(self, url: str) -> Optional[dict]:
        soup = self.get_page(url)
        if not soup:
            return None
        
        title = soup.select_one('.app-title')
        location = soup.select_one('.location')
        content = soup.select_one('#content')
        
        return {
            "title": title.text.strip() if title else "Unknown",
            "location": location.text.strip() if location else "Remote",
            "description": content.get_text('\n').strip() if content else "",
        }
    
    def get_lever_job(self, url: str) -> Optional[dict]:
        soup = self.get_page(url)
        if not soup:
            return None
        
        title = soup.select_one('.posting-headline h2')
        location = soup.select_one('.location')
        content = soup.select_one('.content')
        
        return {
            "title": title.text.strip() if title else "Unknown",
            "location": location.text.strip() if location else "Remote",
            "description": content.get_text('\n').strip() if content else "",
        }
    
    def get_generic_job(self, url: str) -> Optional[dict]:
        soup = self.get_page(url)
        if not soup:
            return None
        
        # Find title
        title = None
        for sel in ['h1', '.job-title', '.posting-title', '[class*="title"]']:
            el = soup.select_one(sel)
            if el:
                title = el.text.strip()
                break
        
        # Find content
        content = None
        for sel in ['main', 'article', '.job-description', '[class*="description"]', '.content']:
            el = soup.select_one(sel)
            if el:
                content = el
                break
        
        if not content:
            content = soup.body
        
        return {
            "title": title or "Unknown Position",
            "location": "Remote",
            "description": content.get_text('\n').strip() if content else "",
        }
    
    def get_job(self, url: str, ats_type: str) -> Optional[dict]:
        """Route to appropriate scraper based on ATS type"""
        if ats_type == 'greenhouse':
            return self.get_greenhouse_job(url)
        elif ats_type == 'lever':
            return self.get_lever_job(url)
        else:
            return self.get_generic_job(url)


# =============================================================================
# AI ANALYSIS
# =============================================================================

def analyze_job(client: anthropic.Anthropic, description: str) -> dict:
    """Analyze job description with Claude"""
    prompt = f"""Analyze this job description. Return JSON only:

{{
    "required_skills": ["skill1", "skill2"],
    "preferred_skills": ["skill1"],
    "keywords": ["keyword1", "keyword2"],
    "remote_type": "fully_remote" | "hybrid" | "onsite" | "unclear",
    "experience_years": "3-5",
    "red_flags": []
}}

Job Description:
{description[:8000]}

JSON only, no other text:"""

    try:
        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=1500,
            messages=[{"role": "user", "content": prompt}]
        )
        text = response.content[0].text
        text = text.replace('```json', '').replace('```', '').strip()
        return json.loads(text)
    except Exception as e:
        logger.error(f"Analysis error: {e}")
        return {}


def calculate_match(analysis: dict, resume: MasterResume) -> float:
    """Calculate match score between job and resume"""
    if not analysis:
        return 0.0
    
    score = 0
    max_score = 0
    
    resume_skills = [s.lower() for s in resume.skills]
    
    for skill in analysis.get('required_skills', []):
        max_score += 2
        if any(skill.lower() in rs or rs in skill.lower() for rs in resume_skills):
            score += 2
    
    for skill in analysis.get('preferred_skills', []):
        max_score += 1
        if any(skill.lower() in rs or rs in skill.lower() for rs in resume_skills):
            score += 1
    
    if analysis.get('remote_type') == 'fully_remote':
        score += 2
        max_score += 2
    
    return score / max_score if max_score > 0 else 0.5


def tailor_resume(client: anthropic.Anthropic, resume: MasterResume, 
                  analysis: dict, job_title: str, company: str) -> dict:
    """Generate tailored resume content"""
    prompt = f"""Tailor this resume for the job. Return JSON only.

Resume:
{json.dumps(asdict(resume), indent=2)}

Job: {job_title} at {company}
Requirements: {json.dumps(analysis.get('required_skills', []))}
Keywords: {json.dumps(analysis.get('keywords', []))}

Return JSON:
{{
    "summary": "2-3 sentence tailored summary",
    "experience": [
        {{"title": "...", "company": "...", "dates": "...", "bullets": ["...", "..."]}}
    ],
    "skills": ["reordered", "skills", "list"]
}}

JSON only:"""

    try:
        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=2500,
            messages=[{"role": "user", "content": prompt}]
        )
        text = response.content[0].text
        text = text.replace('```json', '').replace('```', '').strip()
        return json.loads(text)
    except Exception as e:
        logger.error(f"Resume tailoring error: {e}")
        return {}


def generate_cover_letter(client: anthropic.Anthropic, resume: MasterResume,
                          analysis: dict, job_title: str, company: str) -> str:
    """Generate cover letter"""
    prompt = f"""Write a cover letter (3-4 paragraphs).

Candidate: {resume.name}, {resume.experience[0]['title'] if resume.experience else 'Professional'}
Job: {job_title} at {company}
Key Skills: {', '.join(resume.skills[:8])}
Requirements: {', '.join(analysis.get('required_skills', [])[:5])}

Write the letter (no JSON, just the text):"""

    try:
        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=1200,
            messages=[{"role": "user", "content": prompt}]
        )
        return response.content[0].text.strip()
    except Exception as e:
        logger.error(f"Cover letter error: {e}")
        return ""


# =============================================================================
# DOCUMENT GENERATION
# =============================================================================

def create_resume_docx(content: dict, resume: MasterResume, path: Path):
    """Create tailored resume as .docx"""
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
    
    # Header
    name_p = doc.add_paragraph()
    name_p.add_run(resume.name).bold = True
    name_p.runs[0].font.size = Pt(18)
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    contact = doc.add_paragraph()
    contact.add_run(f"{resume.email} | {resume.phone} | {resume.location}")
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Summary
    doc.add_paragraph()
    doc.add_paragraph().add_run("PROFESSIONAL SUMMARY").bold = True
    doc.add_paragraph(content.get('summary', resume.summary))
    
    # Experience
    doc.add_paragraph()
    doc.add_paragraph().add_run("EXPERIENCE").bold = True
    
    for exp in content.get('experience', resume.experience):
        p = doc.add_paragraph()
        p.add_run(exp['title']).bold = True
        p.add_run(f" | {exp['company']} | {exp['dates']}")
        
        for bullet in exp.get('bullets', []):
            doc.add_paragraph(bullet, style='List Bullet')
    
    # Education
    doc.add_paragraph()
    doc.add_paragraph().add_run("EDUCATION").bold = True
    for edu in resume.education:
        doc.add_paragraph(f"{edu['degree']} - {edu['school']}, {edu.get('year', '')}")
    
    # Skills
    doc.add_paragraph()
    doc.add_paragraph().add_run("SKILLS").bold = True
    skills = content.get('skills', resume.skills)
    doc.add_paragraph(" | ".join(skills))
    
    # Certifications
    if resume.certifications:
        doc.add_paragraph()
        doc.add_paragraph().add_run("CERTIFICATIONS").bold = True
        for cert in resume.certifications:
            doc.add_paragraph(f"• {cert}")
    
    doc.save(path)


def create_cover_letter_docx(letter: str, resume: MasterResume, 
                              job_title: str, company: str, path: Path):
    """Create cover letter as .docx"""
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
    doc.add_paragraph()
    doc.add_paragraph(f"Re: {job_title} Position at {company}")
    doc.add_paragraph()
    
    for para in letter.split('\n\n'):
        if para.strip():
            doc.add_paragraph(para.strip())
    
    doc.add_paragraph()
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph(resume.name)
    doc.add_paragraph(resume.email)
    
    doc.save(path)


# =============================================================================
# MAIN WORKFLOW
# =============================================================================

def load_resume() -> MasterResume:
    path = CONFIG["master_resume_path"]
    if path.exists():
        with open(path) as f:
            return MasterResume(**json.load(f))
    else:
        logger.error(f"Resume not found at {path}")
        sys.exit(1)


def filter_by_title(title: str) -> bool:
    """Check if job title matches our targets"""
    title_lower = title.lower()
    
    # Check exclusions
    if any(exc in title_lower for exc in CONFIG['exclude_keywords']):
        return False
    
    # Check matches
    return any(role in title_lower for role in CONFIG['target_roles'])


def process_job_url(conn: sqlite3.Connection, client: anthropic.Anthropic,
                    scraper: Scraper, resume: MasterResume, 
                    url: str, classification: dict) -> Optional[Job]:
    """Process a single job URL"""
    
    ats_type = classification.get('type', 'unknown')
    company = classification.get('company_slug', 'Unknown')
    
    logger.info(f"Processing: {url}")
    
    # Scrape job details
    details = scraper.get_job(url, ats_type)
    if not details or not details.get('description'):
        logger.warning(f"Could not get job details from {url}")
        return None
    
    # Filter by title
    if not filter_by_title(details['title']):
        logger.info(f"Skipping (title mismatch): {details['title']}")
        return None
    
    # Analyze with AI
    analysis = analyze_job(client, details['description'])
    
    # Calculate match
    match_score = calculate_match(analysis, resume)
    
    if match_score < CONFIG['min_match_score']:
        logger.info(f"Skipping (low match {match_score:.0%}): {details['title']}")
        return None
    
    logger.info(f"Match {match_score:.0%}: {details['title']} at {company}")
    
    # Create job record
    job_id = hashlib.md5(url.encode()).hexdigest()[:12]
    job = Job(
        id=job_id,
        title=details['title'],
        company=company,
        url=url,
        location=details.get('location', 'Remote'),
        description=details['description'],
        requirements=analysis.get('required_skills', []),
        keywords=analysis.get('keywords', []),
        discovered_at=datetime.now().isoformat(),
        match_score=match_score,
        source=ats_type,
    )
    
    # Create output directory
    output_dir = CONFIG['output_dir'] / job_id
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Tailor resume
    tailored = tailor_resume(client, resume, analysis, job.title, job.company)
    resume_path = None
    if tailored:
        resume_path = output_dir / f"resume_{company}_{job_id}.docx"
        create_resume_docx(tailored, resume, resume_path)
        logger.info(f"  Created resume: {resume_path}")
    
    # Generate cover letter
    letter = generate_cover_letter(client, resume, analysis, job.title, job.company)
    cl_path = None
    if letter:
        cl_path = output_dir / f"cover_letter_{company}_{job_id}.docx"
        create_cover_letter_docx(letter, resume, job.title, job.company, cl_path)
        logger.info(f"  Created cover letter: {cl_path}")
    
    # Save analysis
    with open(output_dir / "analysis.json", 'w') as f:
        json.dump({
            "job": asdict(job),
            "analysis": analysis,
            "tailored": tailored,
        }, f, indent=2)
    
    # Save to database
    save_job(conn, job, str(resume_path) if resume_path else None, 
             str(cl_path) if cl_path else None)
    
    logger.info(f"✓ Processed: {job.title} at {job.company}")
    return job


def run_discovery(conn: sqlite3.Connection, queries: list[str] = None):
    """Run Google search discovery"""
    logger.info("Starting job discovery...")
    
    discovery = JobDiscovery()
    results = discovery.discover(queries=queries, max_results_per_query=30)
    
    new_count = 0
    for item in results:
        url = item['url']
        if not url_seen(conn, url):
            save_discovered(conn, url, item.get('classification', {}))
            new_count += 1
    
    logger.info(f"Discovery complete: {new_count} new URLs saved")
    return new_count


def run_processing(conn: sqlite3.Connection, client: anthropic.Anthropic,
                   scraper: Scraper, resume: MasterResume, limit: int = None):
    """Process pending discovered URLs"""
    limit = limit or CONFIG['max_jobs_per_run']
    pending = get_pending_urls(conn, limit)
    
    logger.info(f"Processing {len(pending)} pending URLs...")
    
    processed = 0
    for item in pending:
        url = item['url']
        classification = item['classification']
        
        try:
            job = process_job_url(conn, client, scraper, resume, url, classification)
            if job:
                processed += 1
        except Exception as e:
            logger.error(f"Error processing {url}: {e}")
        finally:
            mark_processed(conn, url)
        
        time.sleep(2)  # Rate limiting
    
    logger.info(f"Processed {processed} jobs")
    return processed


def run_full(discover: bool = True):
    """Run full discovery + processing pipeline"""
    conn = init_db()
    client = anthropic.Anthropic(api_key=CONFIG['anthropic_api_key'])
    scraper = Scraper()
    resume = load_resume()
    
    CONFIG['output_dir'].mkdir(parents=True, exist_ok=True)
    
    if discover:
        run_discovery(conn)
    
    run_processing(conn, client, scraper, resume)
    conn.close()


def show_stats():
    """Show database statistics"""
    conn = init_db()
    cursor = conn.cursor()
    
    print("\n=== Job Hunter Stats ===\n")
    
    cursor.execute("SELECT COUNT(*) FROM jobs")
    print(f"Total jobs processed: {cursor.fetchone()[0]}")
    
    cursor.execute("SELECT COUNT(*) FROM discovered_urls WHERE processed = 0")
    print(f"Pending URLs: {cursor.fetchone()[0]}")
    
    cursor.execute("""
        SELECT title, company, match_score FROM jobs 
        WHERE status = 'new' ORDER BY match_score DESC LIMIT 10
    """)
    
    print("\nTop matches:")
    for row in cursor.fetchall():
        print(f"  [{row[2]:.0%}] {row[0]} at {row[1]}")
    
    conn.close()


def main():
    parser = argparse.ArgumentParser(description="Job Hunter")
    parser.add_argument("--discover-only", action="store_true", help="Only run discovery")
    parser.add_argument("--process-only", action="store_true", help="Only process pending")
    parser.add_argument("--url", type=str, help="Process single URL")
    parser.add_argument("--daemon", action="store_true", help="Run continuously")
    parser.add_argument("--interval", type=int, default=60, help="Daemon interval (minutes)")
    parser.add_argument("--stats", action="store_true", help="Show statistics")
    
    args = parser.parse_args()
    
    if args.stats:
        show_stats()
        return
    
    if args.url:
        conn = init_db()
        client = anthropic.Anthropic(api_key=CONFIG['anthropic_api_key'])
        scraper = Scraper()
        resume = load_resume()
        CONFIG['output_dir'].mkdir(parents=True, exist_ok=True)
        
        classification = URLClassifier.classify(args.url)
        process_job_url(conn, client, scraper, resume, args.url, classification)
        conn.close()
        return
    
    if args.daemon:
        logger.info(f"Starting daemon mode (every {args.interval} min)")
        schedule.every(args.interval).minutes.do(run_full)
        run_full()
        while True:
            schedule.run_pending()
            time.sleep(60)
    elif args.discover_only:
        conn = init_db()
        run_discovery(conn)
        conn.close()
    elif args.process_only:
        conn = init_db()
        client = anthropic.Anthropic(api_key=CONFIG['anthropic_api_key'])
        scraper = Scraper()
        resume = load_resume()
        CONFIG['output_dir'].mkdir(parents=True, exist_ok=True)
        run_processing(conn, client, scraper, resume)
        conn.close()
    else:
        run_full()


if __name__ == "__main__":
    main()
